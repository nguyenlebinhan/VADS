from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import fitz
from docx import Document


def minimal_pdf() -> bytes:
    return b"%PDF-1.7\n1 0 obj\n<<>>\nendobj\ntrailer\n<<>>\n%%EOF\n"


def minimal_docx() -> bytes:
    buffer = BytesIO()
    with ZipFile(buffer, mode="w", compression=ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
        )
        archive.writestr(
            "word/document.xml",
            '<?xml version="1.0"?><w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
        )
    return buffer.getvalue()


def docx_with_table() -> bytes:
    document = Document()
    document.add_paragraph("Điều 1. Phạm vi điều chỉnh")
    document.add_paragraph("Văn bản này quy định việc quản lý hồ sơ hành chính.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Mẫu số"
    table.cell(0, 1).text = "Tên hồ sơ"
    table.cell(1, 0).text = "01"
    table.cell(1, 1).text = "Hồ sơ kiểm tra"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def pdf_with_pages(page_texts: list[str | None]) -> bytes:
    document = fitz.open()
    for text in page_texts:
        page = document.new_page()
        if text:
            y = 72
            for line in text.splitlines():
                page.insert_text((72, y), line, fontsize=11)
                y += 20
    result = document.tobytes()
    document.close()
    return result
