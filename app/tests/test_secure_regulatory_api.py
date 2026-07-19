import json
from datetime import date
from io import BytesIO
from uuid import uuid4

from docx import Document as WordDocument
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session, sessionmaker

from app.core.permissions import UserRole
from app.core.security import hash_password
from app.knowledge_graph.models import GraphVersion, KnowledgeEdge, KnowledgeNode
from app.model.documents import Document
from app.model.processing import ProcessingStatus
from app.model.tenancy import Commune, Province
from app.model.users import User, UserStatus
from app.model.workspaces import Workspace
from app.orchestrator.models import AIWorkflow
from app.regulatory_change.models import (
    RegulatoryDocumentFamily,
    RegulatoryDocumentVersion,
)

PASSWORD = "Correct-Horse-Battery-42!"


def _docx(text: str) -> bytes:
    document = WordDocument()
    for line in text.splitlines():
        document.add_paragraph(line)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _seed_regulatory_document(
    session: Session,
    *,
    commune: Commune,
    owner: User,
    workspace: Workspace,
    title: str,
) -> Document:
    document = Document(
        id=f"doc-{uuid4()}",
        commune_id=commune.id,
        owner_id=owner.id,
        uploaded_by=owner.id,
        workspace_id=workspace.id,
        display_name=title,
        original_filename=f"{title}.pdf",
        mime_type="application/pdf",
        file_extension=".pdf",
        file_size=100,
        checksum=uuid4().hex + uuid4().hex,
        status=ProcessingStatus.COMPLETED,
    )
    family = RegulatoryDocumentFamily(
        workspace_id=workspace.id,
        family_key=f"family-{uuid4()}",
        canonical_title=title,
        document_type="NGHI_DINH",
    )
    session.add_all([document, family])
    session.flush()
    session.add(
        RegulatoryDocumentVersion(
            family_id=family.id,
            document_id=document.id,
            version_number=1,
            title=title,
            document_number="01/2026/ND-CP",
            legal_document_type="NGHI_DINH",
            issuing_agency="Chính phủ",
            issued_date=date(2026, 1, 1),
            effective_date=date(2026, 2, 1),
            domain="Đầu tư công",
            applicable_subjects=["Cơ quan nhà nước"],
            raw_text="Điều 1. Phạm vi áp dụng.",
            executive_summary="Tóm tắt được lưu bởi Regulatory Change API.",
        )
    )
    return document


def _login(client: TestClient, username: str) -> dict[str, str]:
    response = client.post(
        "/api/v1/auth/login",
        json={
            "identifier": username,
            "password": PASSWORD,
            "device_name": "pytest-regulatory",
        },
    )
    assert response.status_code == 200, response.text
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def test_secure_intelligence_routes_require_auth_and_scope_documents_to_owner(
    client: TestClient,
    session_factory: sessionmaker[Session],
) -> None:
    with session_factory() as session:
        province = Province(id=str(uuid4()), name="Tỉnh kiểm thử", code="TEST-RCI")
        commune_a = Commune(
            id=str(uuid4()),
            province_id=province.id,
            name="Xã A",
            code="RCI-A",
        )
        commune_b = Commune(
            id=str(uuid4()),
            province_id=province.id,
            name="Xã B",
            code="RCI-B",
        )
        user_a = User(
            id=str(uuid4()),
            commune_id=commune_a.id,
            username="regulatory.a",
            email="regulatory.a@example.gov.vn",
            full_name="Chuyên viên A",
            role=UserRole.USER,
            password_hash=hash_password(PASSWORD),
            is_active=True,
            must_change_password=False,
            status=UserStatus.ACTIVE,
        )
        admin_a = User(
            id=str(uuid4()),
            commune_id=commune_a.id,
            username="regulatory.admin",
            email="regulatory.admin@example.gov.vn",
            full_name="Quản trị viên A",
            role=UserRole.ADMIN,
            password_hash=hash_password(PASSWORD),
            is_active=True,
            must_change_password=False,
            status=UserStatus.ACTIVE,
        )
        user_b = User(
            id=str(uuid4()),
            commune_id=commune_b.id,
            username="regulatory.b",
            email="regulatory.b@example.gov.vn",
            full_name="Chuyên viên B",
            role=UserRole.USER,
            password_hash=hash_password(PASSWORD),
            is_active=True,
            must_change_password=False,
            status=UserStatus.ACTIVE,
        )
        workspace_a = Workspace(id=str(uuid4()), name="Regulatory A", owner_id=user_a.id)
        workspace_b = Workspace(id=str(uuid4()), name="Regulatory B", owner_id=user_b.id)
        session.add_all(
            [
                province,
                commune_a,
                commune_b,
                user_a,
                admin_a,
                user_b,
                workspace_a,
                workspace_b,
            ]
        )
        session.flush()
        own = _seed_regulatory_document(
            session,
            commune=commune_a,
            owner=user_a,
            workspace=workspace_a,
            title="Văn bản thuộc người dùng A",
        )
        other = _seed_regulatory_document(
            session,
            commune=commune_b,
            owner=user_b,
            workspace=workspace_b,
            title="Văn bản thuộc người dùng B",
        )
        workflow = AIWorkflow(
            id=f"wf-{uuid4()}",
            document_id=own.id,
            intent="KNOWLEDGE_GRAPH_GENERATION",
            status="COMPLETED",
            private_processing=False,
            plan={},
            result={},
        )
        session.add(workflow)
        session.flush()
        graph_version = GraphVersion(
            document_id=own.id,
            workflow_id=workflow.id,
            version=1,
            status="COMPLETED",
            is_current=True,
            model_pipeline=["test-model"],
            validation_issues=[],
        )
        session.add(graph_version)
        session.flush()
        source_node = KnowledgeNode(
            graph_version_id=graph_version.id,
            document_id=own.id,
            source_key="agency-1",
            node_type="AGENCY",
            name="UBND tỉnh",
            canonical_name="UBND tỉnh",
            normalized_key="ubnd tinh",
            properties={"role": "Chủ trì"},
            importance="HIGH",
            confidence=0.95,
        )
        target_node = KnowledgeNode(
            graph_version_id=graph_version.id,
            document_id=own.id,
            source_key="report-1",
            node_type="REPORT",
            name="Báo cáo định kỳ",
            canonical_name="Báo cáo định kỳ",
            normalized_key="bao cao dinh ky",
            properties={},
            importance="MEDIUM",
            confidence=0.9,
        )
        session.add_all([source_node, target_node])
        session.flush()
        session.add(
            KnowledgeEdge(
                graph_version_id=graph_version.id,
                document_id=own.id,
                source_key="submits-1",
                source_node_id=source_node.id,
                target_node_id=target_node.id,
                edge_type="SUBMITS",
                properties={},
                importance="HIGH",
                confidence=0.92,
                verification_status="VERIFIED",
            )
        )
        session.commit()
        own_id = own.id
        other_id = other.id

    assert client.get("/api/v1/regulatory/documents").status_code == 401
    headers = _login(client, "regulatory.a")

    listed = client.get("/api/v1/regulatory/documents", headers=headers)
    assert listed.status_code == 200, listed.text
    assert [item["documentId"] for item in listed.json()["data"]] == [own_id]

    summary = client.get(
        f"/api/v1/regulatory/documents/{own_id}/summary",
        headers=headers,
    )
    assert summary.status_code == 200
    assert summary.json()["data"]["executiveSummary"].startswith("Tóm tắt")

    cross_tenant = client.get(
        f"/api/v1/regulatory/documents/{other_id}/summary",
        headers=headers,
    )
    assert cross_tenant.status_code == 404

    graph = client.get(
        f"/api/v1/documents/{own_id}/knowledge-graph",
        headers=headers,
    )
    assert graph.status_code == 200, graph.text
    graph_data = graph.json()["data"]
    assert {node["name"] for node in graph_data["nodes"]} == {
        "UBND tỉnh",
        "Báo cáo định kỳ",
    }
    assert graph_data["edges"][0]["type"] == "SUBMITS"

    admin_headers = _login(client, "regulatory.admin")
    admin_graph = client.get(
        f"/api/v1/documents/{own_id}/knowledge-graph",
        headers=admin_headers,
    )
    assert admin_graph.status_code == 200, admin_graph.text

    graph_cross_tenant = client.get(
        f"/api/v1/documents/{other_id}/knowledge-graph",
        headers=headers,
    )
    assert graph_cross_tenant.status_code == 404

    upload = client.post(
        "/api/v1/regulatory/documents",
        headers=headers,
        files={
            "file": (
                "quyet-dinh.docx",
                _docx("Điều 1. Quy định mới.\nThời hạn báo cáo: 15 ngày."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            "metadata": (
                None,
                json.dumps(
                    {
                        "familyKey": "quyet-dinh-bao-cao",
                        "title": "Quyết định về chế độ báo cáo",
                        "documentNumber": "02/2026/QD-UBND",
                        "documentType": "QUYET_DINH",
                        "issuingAgency": "UBND tỉnh",
                        "issuedDate": "2026-03-01",
                        "effectiveDate": "2026-04-01",
                        "domain": "Hành chính",
                        "applicableSubjects": ["Cơ quan chuyên môn"],
                    },
                    ensure_ascii=False,
                ),
                "application/json",
            ),
        },
    )
    assert upload.status_code == 201, upload.text
    uploaded = upload.json()["data"]
    assert uploaded["title"] == "Quyết định về chế độ báo cáo"

    listed_after_upload = client.get(
        "/api/v1/regulatory/documents",
        headers=headers,
    )
    visible_ids = {item["documentId"] for item in listed_after_upload.json()["data"]}
    assert visible_ids == {own_id, uploaded["documentId"]}
