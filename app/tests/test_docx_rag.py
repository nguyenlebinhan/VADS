from __future__ import annotations

import os
import tempfile
import unittest
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.docx_rag import embeddings as embeddings_module
from app.docx_rag.chunker import chunk_blocks
from app.docx_rag.docx_reader import read_docx_directory
from app.docx_rag.embeddings import resolve_api_key
from app.docx_rag.index import DocxRagIndex
from app.docx_rag.router import get_docx_rag_service, get_source_store, router
from app.docx_rag.schemas import (
    BlockKind,
    DocxRagResult,
    NoDocxFilesError,
    SourceCitation,
)
from app.docx_rag.service import DocxRagService
from app.docx_rag.source_store import InMemorySourceStore


class FakeOpenAIClient:
    embedding_model = "fake-embedding"

    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        return [[float("năm" in text.casefold()), 1.0] for text in texts]

    def answer_with_context(self, *, question: str, context: str) -> str:
        del question
        assert "clause=Khoản 1" not in context
        return "Thời hạn lưu hồ sơ là năm năm. [Nguồn 1]"


class EmbeddingFailureClient(FakeOpenAIClient):
    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        del texts
        raise RuntimeError("embedding service unavailable")


class FakeDocxRagService:
    def __init__(self) -> None:
        self.answer_calls = 0

    def answer(
        self,
        question: str,
        *,
        top_k: int = 5,
        force_rebuild: bool = False,
    ) -> DocxRagResult:
        del question, top_k, force_rebuild
        self.answer_calls += 1
        call_number = self.answer_calls
        return DocxRagResult(
            answer=f"Câu trả lời {call_number}. [Nguồn 1]",
            sources=[make_source(call_number)],
            retrieval_mode="embedding",
            page_note="DOCX không có số trang cố định.",
        )


def make_source(number: int) -> SourceCitation:
    return SourceCitation(
        file_name=f"policy-{number}.docx",
        chunk_id=f"policy-{number}.docx:chunk-0001",
        paragraph_index=number,
        paragraph_indices=[number],
        quote=f"Trích dẫn cho request {number}.",
        score=0.9,
    )


class DocxRagTest(unittest.TestCase):
    def test_reader_chunker_and_lexical_index_preserve_citations(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            data_dir = Path(temporary_directory)
            path = data_dir / "policy.docx"
            document = Document()
            document.add_paragraph("Điều 3. Phạm vi áp dụng")
            document.add_paragraph("1. Doanh nghiệp phải lưu hồ sơ trong năm năm.")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Thời hạn"
            table.cell(0, 1).text = "Năm năm"
            document.save(path)

            blocks, files = read_docx_directory(data_dir)

            self.assertEqual([path], files)
            self.assertTrue(any(block.kind == BlockKind.TABLE for block in blocks))
            self.assertEqual("Điều 3", blocks[1].article)
            self.assertEqual("Khoản 1", blocks[1].clause)
            self.assertEqual(2, blocks[1].paragraph_index)
            self.assertEqual(1, blocks[2].table_index)
            self.assertIsNone(blocks[2].page_number)

            chunks = chunk_blocks(blocks, max_chars=500)
            index = DocxRagIndex(chunks)
            results = index.search_lexical("thời hạn lưu hồ sơ năm năm", top_k=2)

            self.assertTrue(results)
            self.assertEqual("Điều 3", results[0].chunk.article)
            self.assertEqual("Khoản 1", results[0].chunk.clause)
            self.assertTrue(
                results[0].chunk.paragraph_indices or results[0].chunk.table_indices
            )

    def test_empty_directory_has_clear_error(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            with self.assertRaisesRegex(NoDocxFilesError, "No DOCX files found"):
                read_docx_directory(Path(temporary_directory))

    def test_api_key_can_be_loaded_from_repository_dotenv(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            env_path = Path(temporary_directory) / ".env"
            env_path.write_text("OPENAI_API_KEY=test-key-from-dotenv\n", encoding="utf-8")
            with (
                patch.object(embeddings_module, "ENV_FILE_PATH", env_path),
                patch.dict(
                    os.environ,
                    {"OPENAI_API_KEY": "", "VADS_OPENAI_API_KEY": ""},
                ),
            ):
                embeddings_module._dotenv_values.cache_clear()
                self.assertEqual("test-key-from-dotenv", resolve_api_key())
                embeddings_module._dotenv_values.cache_clear()

    def test_service_uses_embeddings_and_returns_structured_source(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            data_dir = Path(temporary_directory)
            document = Document()
            document.add_paragraph("Điều 3. Thời hạn")
            document.add_paragraph("1. Hồ sơ được lưu trong năm năm.")
            document.save(data_dir / "policy.docx")
            service = DocxRagService(
                data_dir,
                cache_path=data_dir / "index.json",
                client=FakeOpenAIClient(),  # type: ignore[arg-type]
            )

            result = service.answer("Hồ sơ lưu bao nhiêu năm?", top_k=1)

            self.assertEqual("embedding", result.retrieval_mode)
            self.assertEqual("Thời hạn lưu hồ sơ là năm năm.", result.answer)
            self.assertEqual("policy.docx", result.sources[0].file_name)
            self.assertEqual("Điều 3", result.sources[0].article)
            self.assertEqual("Mục 1", result.sources[0].clause)
            self.assertIsNone(result.sources[0].page_number)

    def test_service_falls_back_to_lexical_when_embeddings_fail(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            data_dir = Path(temporary_directory)
            document = Document()
            document.add_paragraph("Điều 4. Hồ sơ được lưu trong năm năm.")
            document.save(data_dir / "policy.docx")
            service = DocxRagService(
                data_dir,
                cache_path=data_dir / "index.json",
                client=EmbeddingFailureClient(),  # type: ignore[arg-type]
            )

            result = service.answer("Hồ sơ lưu bao nhiêu năm?", top_k=1)

            self.assertEqual("lexical", result.retrieval_mode)
            self.assertIn("embedding service unavailable", result.embedding_error or "")


class DocxRagApiTest(unittest.TestCase):
    def setUp(self) -> None:
        self.service = FakeDocxRagService()
        self.store = InMemorySourceStore(ttl_seconds=30)
        application = FastAPI()
        application.include_router(router, prefix="/api")
        application.dependency_overrides[get_docx_rag_service] = lambda: self.service
        application.dependency_overrides[get_source_store] = lambda: self.store
        self.client = TestClient(application)

    def tearDown(self) -> None:
        self.client.close()

    def test_repeated_queries_get_distinct_ids_and_isolated_sources(self) -> None:
        request = {
            "question": "Đối tượng áp dụng là ai?",
            "top_k": 5,
            "rebuild_index": False,
        }

        first = self.client.post("/api/docx-rag/query", json=request)
        second = self.client.post("/api/docx-rag/query", json=request)

        self.assertEqual(200, first.status_code)
        self.assertEqual(200, second.status_code)
        first_body = first.json()
        second_body = second.json()
        self.assertNotEqual(first_body["query_id"], second_body["query_id"])
        self.assertNotIn("sources", first_body)
        self.assertNotIn("sources", second_body)
        self.assertTrue(first_body["sources_available"])
        self.assertEqual(1, first_body["source_count"])

        first_sources = self.client.get(
            f"/api/docx-rag/queries/{first_body['query_id']}/sources"
        )
        second_sources = self.client.get(
            f"/api/docx-rag/queries/{second_body['query_id']}/sources"
        )

        self.assertEqual("policy-1.docx", first_sources.json()["sources"][0]["file_name"])
        self.assertEqual("policy-2.docx", second_sources.json()["sources"][0]["file_name"])
        self.assertEqual(first_body["query_id"], first_sources.json()["query_id"])
        self.assertEqual(second_body["query_id"], second_sources.json()["query_id"])
        self.assertEqual(2, self.service.answer_calls)

    def test_unknown_query_id_returns_404_without_running_service(self) -> None:
        response = self.client.get("/api/docx-rag/queries/missing-query/sources")

        self.assertEqual(404, response.status_code)
        self.assertEqual(0, self.service.answer_calls)

    def test_expired_query_id_returns_404(self) -> None:
        current_time = [100.0]
        expiring_store = InMemorySourceStore(
            ttl_seconds=10,
            clock=lambda: current_time[0],
        )
        self.client.app.dependency_overrides[get_source_store] = lambda: expiring_store
        query_response = self.client.post(
            "/api/docx-rag/query",
            json={"question": "Thời hạn là bao lâu?"},
        )
        query_id = query_response.json()["query_id"]
        current_time[0] = 110.0

        response = self.client.get(f"/api/docx-rag/queries/{query_id}/sources")

        self.assertEqual(404, response.status_code)
        self.assertEqual(1, self.service.answer_calls)


class InMemorySourceStoreTest(unittest.TestCase):
    def test_ttl_uses_default_and_environment_configuration(self) -> None:
        with patch.dict(os.environ, {}, clear=True):
            self.assertEqual(1800, InMemorySourceStore().ttl_seconds)
        with patch.dict(os.environ, {"VADS_DOCX_RAG_QUERY_TTL_SECONDS": "45"}):
            self.assertEqual(45, InMemorySourceStore().ttl_seconds)

    def test_ttl_cleanup_delete_and_result_isolation(self) -> None:
        current_time = [0.0]
        store = InMemorySourceStore(ttl_seconds=10, clock=lambda: current_time[0])
        original_source = make_source(1)
        store.save("first", [original_source], "page note")
        store.save("second", [make_source(2)], "page note")

        fetched = store.get("first")
        self.assertIsNotNone(fetched)
        assert fetched is not None
        fetched.sources[0].file_name = "mutated.docx"
        self.assertEqual("policy-1.docx", store.get("first").sources[0].file_name)  # type: ignore[union-attr]

        store.delete("second")
        self.assertIsNone(store.get("second"))
        current_time[0] = 10.0
        self.assertEqual(1, store.cleanup_expired())
        self.assertIsNone(store.get("first"))

    def test_parallel_writes_keep_sources_under_their_query_ids(self) -> None:
        store = InMemorySourceStore(ttl_seconds=30)

        def save_source(number: int) -> None:
            store.save(str(number), [make_source(number)], "page note")

        with ThreadPoolExecutor(max_workers=8) as executor:
            list(executor.map(save_source, range(50)))

        for number in range(50):
            stored = store.get(str(number))
            self.assertIsNotNone(stored)
            assert stored is not None
            self.assertEqual(f"policy-{number}.docx", stored.sources[0].file_name)


if __name__ == "__main__":
    unittest.main()
