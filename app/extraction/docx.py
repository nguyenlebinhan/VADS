from io import BytesIO

from docx import Document as WordDocument

from app.exceptions import AppError
from app.extraction.types import ExtractedBlock, ExtractedDocument, ExtractedPage, ExtractedTable
from app.model.documents import DocumentType


class DocxExtractor:
    """Extract DOCX logical content; Word pagination is intentionally not fabricated."""

    def extract(self, content: bytes) -> ExtractedDocument:
        try:
            document = WordDocument(BytesIO(content))
        except (KeyError, ValueError, OSError) as exc:
            raise AppError(
                status_code=422,
                code="INVALID_DOCX",
                message="Không thể đọc nội dung DOCX.",
            ) from exc

        blocks: list[ExtractedBlock] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            blocks.append(
                ExtractedBlock(
                    text=text,
                    bbox={"x1": 0, "y1": 0, "x2": 0, "y2": 0},
                    order_index=len(blocks),
                    source="DOCX",
                )
            )

        tables: list[ExtractedTable] = []
        for table_index, table in enumerate(document.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            tables.append(
                ExtractedTable(
                    page_start=0,
                    page_end=0,
                    order_index=table_index,
                    header_rows=rows[:1],
                    rows=rows[1:],
                )
            )
            blocks.append(
                ExtractedBlock(
                    text="\n".join(" | ".join(row) for row in rows),
                    bbox={"x1": 0, "y1": 0, "x2": 0, "y2": 0},
                    order_index=len(blocks),
                    block_type="TABLE",
                    source="DOCX",
                )
            )

        page = ExtractedPage(
            page_index=0,
            width=0,
            height=0,
            rotation=0,
            has_text_layer=True,
            image_only=False,
            needs_ocr=False,
            text="\n".join(block.text for block in blocks),
            blocks=blocks,
        )
        return ExtractedDocument(
            document_type=DocumentType.DOCX.value,
            pages=[page],
            tables=tables,
        )
